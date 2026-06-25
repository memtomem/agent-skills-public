"""Core library for parsing messy Word documents (.docx) into Markdown + JSON.

A .docx is structured OOXML, but real documents hide content where a naive
`"\n".join(p.text for p in doc.paragraphs)` never looks: tables interleaved with
text (python-docx exposes paragraphs and tables as *separate* lists, losing
order), tables nested inside table cells, cells merged with gridSpan/vMerge,
bullet/numbered lists, text boxes (`w:txbxContent`, invisible to python-docx),
headers/footers, tracked changes, and comments.

So this library walks the document **in body order** and emits a structured
element tree — the same shape as pdf-parser / xlsx-parser — plus a per-section
confidence score with flags so risky documents (tracked changes pending, text
boxes, nested/merged tables) surface for review.

Requires python-docx. markitdown is an optional cross-check only.
"""

from __future__ import annotations

import json
import os
import re
from dataclasses import dataclass, field, asdict
from typing import Any, Optional

import docx
from docx.document import Document as _Document
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


class DocxError(Exception):
    """User-facing error for an unreadable/missing document."""


@dataclass
class Element:
    type: str                       # heading | paragraph | list_item | table | textbox | note | image
    text: str = ""
    level: Optional[int] = None
    rows: Optional[list[list[str]]] = None
    spans: Optional[list[dict]] = None
    nested: Optional[list[dict]] = None   # table: nested tables by parent-cell position
    note: Optional[str] = None

    def to_dict(self) -> dict:
        d = asdict(self)
        return {k: v for k, v in d.items() if v not in (None, [], "")}


# --------------------------------------------------------------------------
# Body-order iteration (the correctness keystone)
# --------------------------------------------------------------------------

def iter_block_items(parent):
    """Yield Paragraph and Table objects in true document order.

    python-docx's `doc.paragraphs` and `doc.tables` are separate sequences, so
    using them loses the interleaving of text and tables. Walking the body XML
    children preserves it — the Word analogue of reading order.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        parent_elm = parent
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


# --------------------------------------------------------------------------
# Paragraph classification
# --------------------------------------------------------------------------

def _is_list_paragraph(p: Paragraph) -> bool:
    numpr = p._p.find(qn("w:pPr"))
    if numpr is not None and numpr.find(qn("w:numPr")) is not None:
        return True
    return (p.style and p.style.name or "").startswith("List")


def _heading_level(p: Paragraph) -> Optional[int]:
    name = (p.style.name if p.style else "") or ""
    m = re.match(r"Heading (\d+)", name)
    if m:
        return min(int(m.group(1)), 6)
    if name == "Title":
        return 1
    return None


def _para_textboxes(p: Paragraph) -> list[str]:
    """Extract text from any text boxes anchored in this paragraph.

    Text boxes live in `w:txbxContent` (inside drawing / VML), which python-docx
    does not surface — their text would otherwise be silently dropped."""
    out = []
    for tx in p._p.iter(qn("w:txbxContent")):
        txt = "".join(node.text or "" for node in tx.iter(qn("w:t")))
        if txt.strip():
            out.append(re.sub(r"\s+", " ", txt).strip())
    return out


# --------------------------------------------------------------------------
# Tables (nested + merged-cell aware)
# --------------------------------------------------------------------------

def _cell_text(cell: _Cell) -> str:
    """The cell's own paragraph text (nested tables are captured separately, as
    structured data, by parse_table — never flattened into this string)."""
    parts = []
    for block in iter_block_items(cell):
        if isinstance(block, Paragraph) and block.text.strip():
            parts.append(block.text.strip())
    return " / ".join(parts)


def parse_table(table: Table) -> tuple[list[list[str]], list[dict], list[dict]]:
    """Return (rows, spans, nested).

    Handles horizontal merges (w:gridSpan) and vertical merges (w:vMerge) by
    emitting spans on the anchor and blanks on the followers, so the HTML render
    is faithful. ``nested`` lists tables found inside cells as structured data
    ``{row, col, rows, spans, nested}`` (recursive) keyed by parent-cell position
    — never flattened into a cell string, so user-authored text stays escapable
    and the inner rows are preserved.
    """
    rows: list[list[str]] = []
    spans: list[dict] = []
    nested: list[dict] = []
    # vmerge anchor tracking: map column-index -> anchor (row_idx) currently open
    vmerge_anchor: dict[int, int] = {}
    vmerge_rowspan_pos: dict[int, int] = {}  # column -> index into spans list

    for ri, row in enumerate(table.rows):
        row_vals: list[str] = []
        col = 0
        for tc in row._tr.tc_lst:
            tcPr = tc.find(qn("w:tcPr"))
            gridspan = 1
            vmerge = None
            if tcPr is not None:
                gs = tcPr.find(qn("w:gridSpan"))
                if gs is not None:
                    gridspan = int(gs.get(qn("w:val")) or 1)
                vm = tcPr.find(qn("w:vMerge"))
                if vm is not None:
                    vmerge = vm.get(qn("w:val")) or "continue"
            cell = _Cell(tc, table)
            for nt in cell.tables:
                nrows, nspans, nnested = parse_table(nt)
                if any(any(r) for r in nrows):
                    nested.append({"row": ri, "col": col, "rows": nrows,
                                   "spans": nspans, "nested": nnested})
            text = _cell_text(cell)

            if vmerge == "continue":
                # part of a vertical merge started above; extend that anchor.
                anchor_col = col
                if anchor_col in vmerge_rowspan_pos:
                    spans[vmerge_rowspan_pos[anchor_col]]["rowspan"] += 1
                row_vals.append("")  # follower blank
            else:
                row_vals.append(text)
                span = {}
                if gridspan > 1:
                    span = {"row": ri, "col": col, "rowspan": 1, "colspan": gridspan}
                if vmerge == "restart":
                    span = span or {"row": ri, "col": col, "rowspan": 1, "colspan": gridspan}
                    spans.append(span)
                    vmerge_rowspan_pos[col] = len(spans) - 1
                elif span:
                    spans.append(span)
            # pad followers for horizontal span
            for _ in range(gridspan - 1):
                row_vals.append("")
            col += gridspan
        rows.append(row_vals)
    return rows, spans, nested


# --------------------------------------------------------------------------
# Document-level signals
# --------------------------------------------------------------------------

def _has_tracked_changes(document) -> bool:
    body = document.element.body
    return body.find(qn("w:ins")) is not None or body.find(qn("w:del")) is not None or \
        any(True for _ in body.iter(qn("w:ins"))) or any(True for _ in body.iter(qn("w:del")))


def _count_images(document) -> int:
    return sum(1 for _ in document.element.body.iter(qn("a:blip")))


def _has_comments(document) -> bool:
    try:
        for rel in document.part.rels.values():
            if "comments" in rel.reltype:
                return True
    except Exception:
        pass
    return False


# --------------------------------------------------------------------------
# Assembly
# --------------------------------------------------------------------------

CONFIDENCE_VERIFY_BELOW = 0.75


def build_document(path: str) -> dict:
    if not os.path.exists(path):
        raise DocxError(f"file not found: {path!r}")
    try:
        document = docx.Document(path)
    except Exception as e:
        raise DocxError(f"cannot open {path!r}: {e}") from e

    elements: list[Element] = []
    flags: list[str] = []

    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            for tb in _para_textboxes(block):
                elements.append(Element(type="textbox", text=tb,
                                        note="text recovered from a text box"))
                if "textbox-content" not in flags:
                    flags.append("textbox-content")
            text = block.text.strip()
            if not text:
                continue
            lvl = _heading_level(block)
            if lvl:
                elements.append(Element(type="heading", text=text, level=lvl))
            elif _is_list_paragraph(block):
                elements.append(Element(type="list_item", text=text))
            else:
                elements.append(Element(type="paragraph", text=text))
        elif isinstance(block, Table):
            rows, spans, nested = parse_table(block)
            if not any(any(r) for r in rows):
                continue
            el = Element(type="table", rows=rows)
            notes = []
            if spans:
                el.spans = spans
                if "merged-cells" not in flags:
                    flags.append("merged-cells")
            if nested:
                el.nested = nested
                notes.append("nested-table")
                if "nested-table" not in flags:
                    flags.append("nested-table")
            if notes:
                el.note = "; ".join(notes)
            elements.append(el)

    # headers/footers (first section, deduped)
    seen_hf = set()
    for section in document.sections:
        for part, label in ((section.header, "header"), (section.footer, "footer")):
            try:
                txt = "\n".join(p.text for p in part.paragraphs if p.text.strip()).strip()
            except Exception:
                txt = ""
            if txt and txt not in seen_hf:
                seen_hf.add(txt)
                elements.append(Element(type="note", text=txt, note=label))

    # document-level signals
    TEXT_TYPES = {"heading", "paragraph", "list_item", "table", "textbox"}
    text_bearing = sum(1 for e in elements if e.type in TEXT_TYPES)
    n_img = _count_images(document)
    # Image-dominant document (embedded images, no extractable text) carries
    # content local tools can't read — route it to a vision pass explicitly
    # (parity with pdf-parser / pptx-parser needs_vision), not just a flag.
    needs_vision = bool(n_img) and text_bearing == 0
    if n_img:
        flags.append("embedded-images-not-extracted")
        note = ("image-only document — render and transcribe with a vision model"
                if needs_vision else "images are referenced but not extracted in this pass")
        elements.append(Element(type="image", text=f"{n_img} embedded image(s)", note=note))
    if needs_vision:
        flags.append("needs-vision")
    if _has_tracked_changes(document):
        flags.append("tracked-changes-present")
    if _has_comments(document):
        flags.append("comments-present")

    confidence = 1.0
    caps = {
        "tracked-changes-present": 0.65,
        "textbox-content": 0.80,
        "nested-table": 0.75,
        "merged-cells": 0.80,
        "comments-present": 0.85,
        "embedded-images-not-extracted": 0.85,
        "needs-vision": 0.40,
    }
    for f in flags:
        confidence = min(confidence, caps.get(f, 1.0))

    return {
        "source": os.path.basename(path),
        "element_count": len(elements),
        "confidence": round(confidence, 2),
        "flags": flags,
        "needs_review": round(confidence, 2) < CONFIDENCE_VERIFY_BELOW,
        "needs_vision": needs_vision,
        "elements": [e.to_dict() for e in elements],
    }


# --------------------------------------------------------------------------
# Rendering
# --------------------------------------------------------------------------

def _esc(s: str) -> str:
    return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _rows_to_gfm(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    norm = [r + [""] * (width - len(r)) for r in rows]
    def fmt(r):
        # escape HTML-sensitive chars too, not just pipes — a GFM cell still ends
        # up in rendered HTML, so raw user "<script>" must not survive.
        return "| " + " | ".join(_esc(c).replace("|", "\\|").replace("\n", " ") for c in r) + " |"
    out = [fmt(norm[0]), "| " + " | ".join(["---"] * width) + " |"]
    out += [fmt(r) for r in norm[1:]]
    return "\n".join(out)


def _rows_to_html(rows: list[list[str]], spans: Optional[list[dict]],
                  nested: Optional[list[dict]] = None) -> str:
    spans = spans or []
    nested = nested or []
    anchor = {(s["row"], s["col"]): (s.get("rowspan", 1), s.get("colspan", 1)) for s in spans}
    nested_at: dict[tuple, list] = {}
    for nt in nested:
        nested_at.setdefault((nt["row"], nt["col"]), []).append(nt)
    covered = set()
    for s in spans:
        for dr in range(s.get("rowspan", 1)):
            for dc in range(s.get("colspan", 1)):
                if dr or dc:
                    covered.add((s["row"] + dr, s["col"] + dc))
    out = ["<table>"]
    for ri, row in enumerate(rows):
        tag = "th" if ri == 0 else "td"
        cells = []
        for ci, val in enumerate(row):
            if (ri, ci) in covered:
                continue
            attr = ""
            if (ri, ci) in anchor:
                rs, cs = anchor[(ri, ci)]
                if rs > 1:
                    attr += f' rowspan="{rs}"'
                if cs > 1:
                    attr += f' colspan="{cs}"'
            # ALL cell text is escaped — never trusted as HTML. A nested table is
            # rendered by this same (escaping) function and appended, so its inner
            # text is safe too; user-authored "<table>"/"<script>" text cannot
            # escape into raw markup.
            content = _esc(val).replace(chr(10), " ")
            for nt in nested_at.get((ri, ci), []):
                content += _rows_to_html(nt["rows"], nt.get("spans"), nt.get("nested"))
            cells.append(f"<{tag}{attr}>{content}</{tag}>")
        out.append("  <tr>" + "".join(cells) + "</tr>")
    out.append("</table>")
    return "\n".join(out)


def element_to_markdown(el: dict) -> str:
    # User-authored text is HTML-escaped on the way into Markdown — the rendered
    # output may be turned back into HTML downstream, so raw "<script>" from a
    # source document must never pass through. (JSON keeps the raw text.)
    t = el.get("type")
    text = _esc((el.get("text") or "").strip())
    if t == "heading":
        return f"{'#' * (el.get('level') or 2)} {text}"
    if t == "paragraph":
        return text
    if t == "list_item":
        return f"- {text}"
    if t == "textbox":
        return f"> 📦 **[text box]** {text}"
    if t == "note":
        return f"<!-- {el.get('note', 'note')} --> {text}"
    if t == "image":
        return f"<!-- 🖼 {el.get('text', '')}: {el.get('note', '')} -->"
    if t == "table":
        rows = el.get("rows") or []
        spans = el.get("spans")
        nested = el.get("nested")
        # A flat pipe table can't carry merges or a nested table — use HTML then.
        body = (_rows_to_html(rows, spans, nested) if (spans or nested)
                else _rows_to_gfm(rows))
        return (f"<!-- table{(' · ' + el['note']) if el.get('note') else ''} -->\n" + body)
    return el.get("text", "").strip()


def to_markdown(document: dict) -> str:
    out = [f"<!-- parsed from {_esc(document.get('source', ''))} -->\n"]
    conf = document.get("confidence")
    if conf is not None and conf < CONFIDENCE_VERIFY_BELOW:
        out.append(f"> 🔎 Low parse confidence ({conf}); verify. "
                   f"flags: {', '.join(document.get('flags') or []) or 'none'}")
    for el in document.get("elements", []):
        md = element_to_markdown(el)
        if md:
            out.append(md)
    return "\n\n".join(out).strip() + "\n"


def save_outputs(document: dict, md_path: str, json_path: str) -> None:
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(document, f, ensure_ascii=False, indent=2)
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(to_markdown(document))


def markitdown_crosscheck(path: str) -> Optional[str]:
    try:
        from markitdown import MarkItDown
    except Exception:
        return None
    try:
        return MarkItDown().convert(path).text_content
    except Exception:
        return None
