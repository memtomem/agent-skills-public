"""Generate messy .docx fixtures (never committed; built on demand).

messy.docx exercises: headings, body text interleaved with a table, a
bullet list, a merged-cell table (horizontal + vertical), a nested table, a
text box, and a header — KR+EN.

image_only.docx: a document whose only content is an embedded image (no text)
— the needs_vision routing case.
"""
import os
import sys
import tempfile

import docx
from docx.oxml.ns import qn
from docx.shared import Pt
from PIL import Image


def _add_merged_table(document):
    # 3 cols x 3 rows; merge top row's last two cells (colspan=2), and
    # vertically merge the first column of rows 2-3 (rowspan=2).
    t = document.add_table(rows=3, cols=3)
    t.style = "Table Grid"
    t.cell(0, 0).text = "Region"
    a = t.cell(0, 1); a.text = "H1"
    t.cell(0, 1).merge(t.cell(0, 2))           # horizontal merge -> colspan 2
    t.cell(1, 0).text = "Asia"
    t.cell(1, 0).merge(t.cell(2, 0))           # vertical merge -> rowspan 2
    t.cell(1, 1).text = "120"; t.cell(1, 2).text = "135"
    t.cell(2, 1).text = "90"; t.cell(2, 2).text = "110"


def _add_nested_table(document):
    outer = document.add_table(rows=1, cols=2)
    outer.style = "Table Grid"
    outer.cell(0, 0).text = "Outer A"
    c = outer.cell(0, 1)
    c.paragraphs[0].text = "inner:"        # text BEFORE the nested table
    inner = c.add_table(rows=2, cols=2)
    inner.cell(0, 0).text = "n1"; inner.cell(0, 1).text = "n2"
    inner.cell(1, 0).text = "n3"; inner.cell(1, 1).text = "n4"


def _add_textbox(document):
    # Inject a minimal w:txbxContent via a run's XML so the parser must recover it.
    p = document.add_paragraph()
    run = p.add_run()
    xml = (
        '<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:v="urn:schemas-microsoft-com:vml">'
        '<v:shape><v:textbox><w:txbxContent>'
        '<w:p><w:r><w:t>중요 공지: 본 문서는 기밀입니다 (Confidential)</w:t></w:r></w:p>'
        '</w:txbxContent></v:textbox></v:shape></w:pict>'
    )
    from docx.oxml import parse_xml
    run._r.append(parse_xml(xml))


def build(path: str) -> str:
    document = docx.Document()
    document.sections[0].header.paragraphs[0].text = "ACME Corp — 2025 보고서"

    document.add_heading("1. 개요 Overview", level=1)
    document.add_paragraph("본 문서는 비정형 워드 파싱을 검증하기 위한 예제이다. "
                           "This document interleaves text, tables, and lists.")
    document.add_heading("2. 재무 Financials", level=2)
    document.add_paragraph("아래 표는 본문 사이에 끼어 있다 (interleaved table):")

    t = document.add_table(rows=3, cols=3)
    t.style = "Table Grid"
    hdr = ["항목 Item", "2024", "2025"]
    for j, v in enumerate(hdr):
        t.cell(0, j).text = v
    for i, (name, a, b) in enumerate([("매출 Revenue", "1050", "1239"),
                                      ("순이익 Net", "70", "112")]):
        t.cell(i + 1, 0).text = name
        t.cell(i + 1, 1).text = a
        t.cell(i + 1, 2).text = b

    document.add_paragraph("핵심 사항 Key points:")
    for item in ["디지털 전환 가속", "해외 시장 확대", "R&D 투자 증대"]:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("3. 병합/중첩 표 Merged & Nested", level=2)
    _add_merged_table(document)
    document.add_paragraph("중첩 표 nested:")
    _add_nested_table(document)

    _add_textbox(document)

    document.save(path)
    return path


def build_image_only(path: str) -> str:
    document = docx.Document()
    fd, tmp = tempfile.mkstemp(suffix=".png"); os.close(fd)
    Image.new("RGB", (240, 160), "teal").save(tmp)
    document.add_picture(tmp)          # one inline image, no text at all
    document.save(path)
    try:
        os.remove(tmp)
    except OSError:
        pass
    return path


def make_all(out_dir: str) -> None:
    os.makedirs(out_dir, exist_ok=True)
    build(os.path.join(out_dir, "messy.docx"))
    build_image_only(os.path.join(out_dir, "image_only.docx"))


if __name__ == "__main__":
    out = sys.argv[1] if len(sys.argv) > 1 else os.path.dirname(__file__)
    make_all(out)
    print("fixtures written to", out)
